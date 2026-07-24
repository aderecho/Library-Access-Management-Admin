from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "UP_Cebu_RFID_Admin_Deployment_Guide.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 55, 72)
MUTED = RGBColor(89, 101, 113)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = RGBColor(255, 255, 255)
CURRENT_STEP_NUM_ID = None


def set_font(run, name="Calibri", size=11, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field(paragraph, field):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, end])
    set_font(run, size=9, color=MUTED)


def add_code(doc, text):
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = False
    run = p.add_run(text)
    set_font(run, name="Menlo", size=9, color=INK)
    return p


def add_note(doc, label, text, warning=False):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF4CE" if warning else LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_font(r, bold=True, color=INK)
    r = p.add_run(text)
    set_font(r, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def start_step_sequence(doc):
    global CURRENT_STEP_NUM_ID
    numbering = doc.part.numbering_part.element
    abstract_nums = numbering.findall(qn("w:abstractNum"))
    nums = numbering.findall(qn("w:num"))
    new_abstract_id = max([int(n.get(qn("w:abstractNumId"))) for n in abstract_nums] + [-1]) + 1
    new_num_id = max([int(n.get(qn("w:numId"))) for n in nums] + [0]) + 1

    list_style = doc.styles["List Number"]._element
    num_pr = list_style.pPr.find(qn("w:numPr"))
    base_num_id = int(num_pr.find(qn("w:numId")).get(qn("w:val")))
    base_num = next(n for n in nums if int(n.get(qn("w:numId"))) == base_num_id)
    base_abstract_id = int(base_num.find(qn("w:abstractNumId")).get(qn("w:val")))
    base_abstract = next(
        n for n in abstract_nums if int(n.get(qn("w:abstractNumId"))) == base_abstract_id
    )

    abstract_copy = deepcopy(base_abstract)
    abstract_copy.set(qn("w:abstractNumId"), str(new_abstract_id))
    numbering.insert(len(abstract_nums), abstract_copy)

    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(new_abstract_id))
    new_num.append(abstract_ref)
    numbering.append(new_num)
    CURRENT_STEP_NUM_ID = new_num_id


def add_step(doc, title, body=None):
    p = doc.add_paragraph(style="List Number")
    if CURRENT_STEP_NUM_ID is None:
        start_step_sequence(doc)
    p_pr = p._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id = num_pr.find(qn("w:numId"))
    if num_id is None:
        num_id = OxmlElement("w:numId")
        num_pr.append(num_id)
    num_id.set(qn("w:val"), str(CURRENT_STEP_NUM_ID))
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_font(r, bold=True, color=DARK_BLUE)
    if body:
        r = p.add_run(f" — {body}")
        set_font(r, color=INK)
    return p


def add_check(doc, text):
    p = doc.add_paragraph(style="Checklist")
    p.add_run(text)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header)
        set_font(r, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, size=10.5)
            if idx == 0:
                set_font(r, size=10.5, bold=True, color=DARK_BLUE)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for name in ("List Bullet", "List Bullet 2", "List Number"):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

checklist = styles.add_style("Checklist", 1)
checklist.font.name = "Calibri"
checklist.font.size = Pt(11)
checklist.paragraph_format.left_indent = Inches(0.375)
checklist.paragraph_format.first_line_indent = Inches(-0.188)
checklist.paragraph_format.space_after = Pt(4)

code_style = styles.add_style("Code Block", 1)
code_style.font.name = "Menlo"
code_style.font.size = Pt(9)
code_style.paragraph_format.left_indent = Inches(0.18)
code_style.paragraph_format.right_indent = Inches(0.18)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(8)
code_style.paragraph_format.line_spacing = 1.15
code_ppr = code_style._element.get_or_add_pPr()
shd = OxmlElement("w:shd")
shd.set(qn("w:fill"), LIGHT_GRAY)
code_ppr.append(shd)

header = section.header
hp = header.paragraphs[0]
hp.text = "UP Cebu RFID Admin  |  Deployment Guide"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(hp.runs[0], size=9, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = fp.add_run("Page ")
set_font(r, size=9, color=MUTED)
add_field(fp, "PAGE")

# Editorial-cover opening.
for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DEPLOYMENT GUIDE")
set_font(r, size=11, bold=True, color=BLUE)
p.paragraph_format.space_after = Pt(16)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("UP Cebu RFID Admin")
set_font(r, size=28, bold=True, color=INK)
p.paragraph_format.space_after = Pt(8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Step-by-step local deployment based on start.sh")
set_font(r, size=15, color=DARK_BLUE)
p.paragraph_format.space_after = Pt(20)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Platform: macOS  •  Laravel 12  •  PHP 8.2+  •  PostgreSQL 14  •  Redis")
set_font(r, size=10.5, color=MUTED)
p.paragraph_format.space_after = Pt(6)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prepared from the repository configuration and start.sh")
set_font(r, size=10.5, italic=True, color=MUTED)

doc.add_page_break()

doc.add_heading("1. Purpose and deployment scope", level=1)
doc.add_paragraph(
    "This guide explains how to prepare, start, verify, stop, and troubleshoot the UP Cebu RFID Admin application by following the repository’s start.sh script. The procedure is intended for a macOS development or on-premises workstation where Homebrew manages PostgreSQL and Redis."
)
add_note(
    doc,
    "Important",
    "start.sh is a development deployment launcher. It runs Laravel’s built-in server, Vite’s development server, live log streaming, a queue worker, Laravel Reverb, and the RFID Redis synchronization process in foreground terminals. It is not a production web-server, process-supervisor, TLS, backup, or high-availability configuration.",
    warning=True,
)

doc.add_heading("2. What start.sh does", level=1)
add_table(
    doc,
    ["Order", "Action", "Result"],
    [
        ("1", "Moves to the script directory", "Commands run from the repository root regardless of the caller’s current folder."),
        ("2", "Checks required commands", "Stops if brew, php, composer, npm, or lsof is unavailable."),
        ("3", "Releases ports 8000, 8080, and 5173", "Stops stale Laravel, Reverb, and Vite listeners before startup."),
        ("4", "Starts Homebrew services", "Starts postgresql@14 and redis."),
        ("5", "Applies migrations", "Runs php artisan migrate --force against the database configured in .env."),
        ("6", "Runs composer run dev", "Starts the six foreground development processes described below."),
    ],
    [900, 2800, 5660],
)

doc.add_heading("3. Runtime processes and ports", level=1)
add_table(
    doc,
    ["Process", "Command", "Default endpoint or role"],
    [
        ("Laravel server", "php artisan serve", "http://localhost:8000"),
        ("Redis queue worker", "php artisan queue:work redis …", "Processes queued jobs; 3 attempts; 60-second timeout."),
        ("Application logs", "php artisan pail --timeout=0", "Streams Laravel logs in the terminal."),
        ("Vite", "npm run dev", "Frontend development server, normally port 5173."),
        ("Reverb", "php artisan reverb:start", "WebSocket server, configured on localhost:8080."),
        ("RFID CDC sync", "php artisan rfid:sync-redis", "Synchronizes RFID transaction read data to Redis."),
    ],
    [1700, 3580, 4080],
)

doc.add_page_break()
doc.add_heading("4. Prerequisites", level=1)
doc.add_paragraph("Complete these checks before the first deployment.")
for text in (
    "A supported macOS workstation with a terminal and administrator access for Homebrew installation.",
    "Homebrew.",
    "PHP 8.2 or newer, including the PostgreSQL PDO extension.",
    "Composer.",
    "Node.js and npm.",
    "PostgreSQL 14 installed as the Homebrew formula postgresql@14.",
    "Redis installed through Homebrew.",
    "lsof, normally included with macOS.",
    "A local copy of this repository.",
):
    add_check(doc, "☐ " + text)

doc.add_heading("4.1 Install the required software", level=2)
doc.add_paragraph("If the tools are not already installed, use Homebrew:")
add_code(doc, "brew install php composer node postgresql@14 redis")
doc.add_paragraph("Confirm that the commands used by start.sh are available:")
add_code(doc, "brew --version\nphp --version\ncomposer --version\nnpm --version\nlsof -v")
add_note(
    doc,
    "PHP check",
    "The application requires PHP ^8.2. If multiple PHP installations exist, confirm that the php executable shown by command -v php is the intended version.",
)

doc.add_heading("4.2 Obtain the source code", level=2)
doc.add_paragraph("Open Terminal and enter the project directory. Replace the example path if the repository is stored elsewhere.")
add_code(doc, 'cd "/path/to/up-cebu-rfid-admin"')

doc.add_heading("5. First-time application setup", level=1)
start_step_sequence(doc)
add_step(doc, "Install PHP dependencies")
add_code(doc, "composer install")
add_step(doc, "Create the environment file")
add_code(doc, "cp .env.example .env")
add_note(doc, "Existing environment", "Do not overwrite .env when it already contains valid deployment settings or secrets.", warning=True)
add_step(doc, "Generate the Laravel application key")
add_code(doc, "php artisan key:generate")
add_step(doc, "Install JavaScript dependencies")
add_code(doc, "npm install")

doc.add_page_break()
doc.add_heading("6. Configure the environment", level=1)
doc.add_paragraph(
    "Edit .env before running start.sh. The current repository deployment uses PostgreSQL, Redis, Reverb, Redis-backed sessions and queues, and the RFID Redis read model. Keep secrets out of version control."
)

doc.add_heading("6.1 PostgreSQL settings", level=2)
doc.add_paragraph("Use settings that match the PostgreSQL instance and database created for this application:")
add_code(
    doc,
    "DB_CONNECTION=pgsql\n"
    "DB_HOST=127.0.0.1\n"
    "DB_PORT=5432\n"
    "DB_DATABASE=library_rfid\n"
    "DB_USERNAME=postgres\n"
    "DB_PASSWORD=<secure-password>",
)
add_note(
    doc,
    "Port consistency",
    "Homebrew PostgreSQL commonly listens on port 5432. If .env uses another port, such as 5433, PostgreSQL must actually be configured to listen there. start.sh starts the service but does not change its port.",
    warning=True,
)

doc.add_heading("6.2 Create the PostgreSQL database", level=2)
doc.add_paragraph("Start PostgreSQL, then create the configured database if it does not already exist:")
add_code(doc, "brew services start postgresql@14\ncreatedb library_rfid")
doc.add_paragraph("Test connectivity using the same host, port, database, and user configured in .env:")
add_code(doc, "psql -h 127.0.0.1 -p 5432 -U postgres -d library_rfid -c 'select 1;'")

doc.add_heading("6.3 Redis, Reverb, and application settings", level=2)
add_code(
    doc,
    "APP_ENV=local\n"
    "APP_DEBUG=true\n"
    "APP_URL=http://localhost:8000\n\n"
    "SESSION_DRIVER=redis\n"
    "QUEUE_CONNECTION=redis\n"
    "CACHE_STORE=redis\n"
    "REDIS_CLIENT=predis\n"
    "REDIS_HOST=127.0.0.1\n"
    "REDIS_PORT=6379\n"
    "RFID_TRANSACTION_READ_MODEL=redis\n\n"
    "BROADCAST_CONNECTION=reverb\n"
    "REVERB_HOST=localhost\n"
    "REVERB_PORT=8080\n"
    "REVERB_SCHEME=http",
)
add_note(
    doc,
    "Credentials",
    "Set unique REVERB_APP_ID, REVERB_APP_KEY, and REVERB_APP_SECRET values. The VITE_REVERB_* values must reference the matching Reverb settings.",
    warning=True,
)

doc.add_heading("6.4 Clear cached configuration after editing .env", level=2)
add_code(doc, "php artisan optimize:clear")

doc.add_page_break()
doc.add_heading("7. Run the deployment", level=1)
doc.add_paragraph("Perform the following steps from the repository root.")
start_step_sequence(doc)
add_step(doc, "Make start.sh executable", "This is normally required only once.")
add_code(doc, "chmod +x start.sh")
add_step(doc, "Start the application stack")
add_code(doc, "./start.sh")
add_step(
    doc,
    "Keep the terminal open",
    "composer run dev owns the foreground processes. Closing the terminal or pressing Control+C stops the application processes.",
)

doc.add_heading("7.1 Expected startup sequence", level=2)
for text in (
    "The script validates brew, php, composer, npm, and lsof.",
    "Any listener on port 8000, 8080, or 5173 is terminated. The script waits up to approximately five seconds for each port.",
    "Homebrew starts PostgreSQL 14 and Redis.",
    "Laravel applies all pending database migrations with --force.",
    "Composer launches Laravel, the Redis queue worker, Pail logs, Vite, Reverb, and the RFID Redis synchronization command.",
):
    add_bullet(doc, text)
add_note(
    doc,
    "Port warning",
    "start.sh terminates any process listening on ports 8000, 8080, or 5173, even if that process belongs to another project. Save work and check those ports before running the script on a shared workstation.",
    warning=True,
)

doc.add_heading("8. Verify the deployment", level=1)
doc.add_paragraph("Use a second terminal so the startup terminal remains active.")
start_step_sequence(doc)
add_step(doc, "Confirm the web application responds")
add_code(doc, "curl -I http://localhost:8000")
doc.add_paragraph("A successful response is normally HTTP 200 or an expected redirect such as HTTP 302.")
add_step(doc, "Open the application")
add_code(doc, "open http://localhost:8000")
add_step(doc, "Confirm PostgreSQL and Redis services")
add_code(doc, "brew services list | grep -E 'postgresql@14|redis'\nredis-cli ping")
doc.add_paragraph("Redis should reply with PONG.")
add_step(doc, "Confirm expected listeners")
add_code(doc, "lsof -nP -iTCP:8000 -sTCP:LISTEN\nlsof -nP -iTCP:8080 -sTCP:LISTEN\nlsof -nP -iTCP:5173 -sTCP:LISTEN")
add_step(doc, "Confirm migration status")
add_code(doc, "php artisan migrate:status")

doc.add_heading("8.1 Deployment acceptance checklist", level=2)
for text in (
    "The login or landing page loads at http://localhost:8000.",
    "No fatal errors appear in the startup terminal.",
    "PostgreSQL and Redis show as started.",
    "redis-cli ping returns PONG.",
    "Ports 8000, 8080, and 5173 have listeners.",
    "All database migrations report Ran.",
    "Real-time updates connect without browser-console WebSocket errors.",
    "Queued work is processed by the queue worker.",
):
    add_check(doc, "☐ " + text)

doc.add_page_break()
doc.add_heading("9. Optional: seed initial or demonstration data", level=1)
doc.add_paragraph(
    "start.sh runs migrations only; it does not run database seeders. Seed only when the target environment is intended to receive the repository’s initial or demonstration records."
)
add_code(doc, "php artisan db:seed")
add_note(
    doc,
    "Data safety",
    "Review DatabaseSeeder and all referenced seeders before running this command against an existing database. Never use migrate:fresh on a database whose data must be preserved.",
    warning=True,
)

doc.add_heading("10. Stop and restart", level=1)
doc.add_heading("10.1 Stop application processes", level=2)
doc.add_paragraph(
    "In the terminal running ./start.sh, press Control+C. The concurrently configuration uses --kill-others, so the managed development processes should stop together."
)
doc.add_heading("10.2 Stop PostgreSQL and Redis when no longer needed", level=2)
add_code(doc, "brew services stop postgresql@14\nbrew services stop redis")
doc.add_heading("10.3 Restart", level=2)
add_code(doc, "./start.sh")

doc.add_heading("11. Updating an existing deployment", level=1)
doc.add_paragraph("After receiving updated source code, use this safe sequence:")
start_step_sequence(doc)
add_step(doc, "Back up important application and database data")
add_step(doc, "Stop the running development stack", "Press Control+C in its terminal.")
add_step(doc, "Update the source code", "Use the team’s approved Git workflow.")
add_step(doc, "Refresh PHP dependencies")
add_code(doc, "composer install")
add_step(doc, "Refresh JavaScript dependencies")
add_code(doc, "npm install")
add_step(doc, "Clear cached Laravel state")
add_code(doc, "php artisan optimize:clear")
add_step(doc, "Start the updated stack", "start.sh applies pending migrations automatically.")
add_code(doc, "./start.sh")
add_step(doc, "Repeat the verification and acceptance checklist")

doc.add_heading("12. Troubleshooting", level=1)
add_table(
    doc,
    ["Symptom", "Likely cause and action"],
    [
        ("Required command not found", "Install the named command, ensure it is on PATH, then rerun ./start.sh."),
        ("Could not release port", "Identify the listener with lsof -nP -iTCP:<port> -sTCP:LISTEN, stop it safely, then retry."),
        ("Database connection refused", "Verify the Homebrew PostgreSQL service, DB_HOST, DB_PORT, database name, user, password, and PostgreSQL listening port."),
        ("Migration fails", "Read the migration error before restarting. Correct connectivity or schema issues; do not delete data as a shortcut."),
        ("Redis connection refused", "Run brew services start redis, then redis-cli ping. Confirm REDIS_HOST and REDIS_PORT."),
        ("Vite does not start", "Run npm install, confirm Node/npm versions, and check whether port 5173 is free."),
        ("Reverb fails or live updates do not work", "Confirm port 8080, matching REVERB/VITE_REVERB settings, and browser WebSocket errors."),
        ("One dev process exits", "concurrently uses --kill-others, so the complete stack may stop. Fix the first failing process, then rerun ./start.sh."),
        ("Permission denied: ./start.sh", "Run chmod +x start.sh once, then retry."),
    ],
    [2500, 6860],
)

doc.add_heading("13. Production deployment boundary", level=1)
doc.add_paragraph(
    "Do not use this start.sh workflow as the final Internet-facing production architecture. A production deployment should separately define:"
)
for text in (
    "A production web server and PHP-FPM configuration.",
    "A process supervisor for queue workers, Reverb, and the RFID synchronization command.",
    "A built frontend asset bundle created with npm run build instead of the Vite development server.",
    "APP_ENV=production, APP_DEBUG=false, a production APP_URL, protected secrets, and restricted Reverb origins.",
    "TLS/HTTPS, firewall rules, authentication hardening, log rotation, monitoring, backups, and tested restore procedures.",
    "A controlled, backed-up migration and rollback procedure.",
):
    add_bullet(doc, text)
add_note(
    doc,
    "Scope",
    "This document intentionally stays faithful to start.sh. Production infrastructure details require the intended server OS, hosting topology, domain/TLS setup, and operations requirements.",
)

doc.add_heading("Appendix A — Command summary", level=1)
add_code(
    doc,
    "# First-time setup\n"
    "composer install\n"
    "cp .env.example .env\n"
    "php artisan key:generate\n"
    "npm install\n"
    "php artisan optimize:clear\n"
    "chmod +x start.sh\n\n"
    "# Deploy/start\n"
    "./start.sh\n\n"
    "# Verify\n"
    "curl -I http://localhost:8000\n"
    "redis-cli ping\n"
    "php artisan migrate:status\n\n"
    "# Stop Homebrew services when finished\n"
    "brew services stop postgresql@14\n"
    "brew services stop redis",
)

doc.core_properties.title = "UP Cebu RFID Admin Deployment Guide"
doc.core_properties.subject = "Step-by-step local deployment based on start.sh"
doc.core_properties.author = "UP Cebu RFID Admin Project"
doc.core_properties.keywords = "Laravel, deployment, start.sh, PostgreSQL, Redis, Reverb, Vite"

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
